"""
Genera el .docx del entregable Gold Standard + Benchmark OCR.

Ejecutar desde la raiz:
    python entregables/_build_docx.py

Produce:
    entregables/Entrega_Gold_Standard_y_Benchmark_OCR.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
FIGS = ROOT / 'figs'
OCR_FIG = ROOT.parent / 'data' / 'processed' / 'fig11_ocr_benchmark.png'
OUT = ROOT / 'Entrega_Gold_Standard_y_Benchmark_OCR.docx'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    return p


def add_para_mixed(doc, parts):
    p = doc.add_paragraph()
    for text, bold, italic, is_code in parts:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if bold: r.bold = True
        if italic: r.italic = True
        if is_code:
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p


def add_table_from_rows(doc, headers, rows, col_widths=None, header_bg='2E5C8A'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = ''
            p = t.rows[i].cells[j].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
    return t


def add_image(doc, path, width_cm=15, caption=None):
    if not Path(path).exists():
        p = doc.add_paragraph(f'[FIGURA FALTANTE: {path}]')
        r = p.runs[0]; r.italic = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═════ PORTADA ═════
title = doc.add_heading('', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Entrega Académica\nGold Standard y Selección del Motor OCR')
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

subt = doc.add_paragraph()
subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subt.add_run('Proyecto SinergIA Lab — IDP para Documentos Corporativos Colombianos SECOP')
r.italic = True
r.font.size = Pt(13)

doc.add_paragraph()

meta = doc.add_table(rows=6, cols=2)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_items = [
    ('Institución', 'Pontificia Universidad Javeriana'),
    ('Programa', 'Especialización en Inteligencia Artificial — Ciclo 1 · 2026'),
    ('Fase CRISP-DM++', '2.1.1 — Selección del Motor OCR'),
    ('Fecha de entrega', '2026-04-20'),
    ('Notebook de referencia', '03_benchmark_ocr.ipynb'),
    ('Bitácora complementaria', 'OCR_BENCHMARK.md'),
]
for i, (k, v) in enumerate(meta_items):
    meta.rows[i].cells[0].text = ''
    p = meta.rows[i].cells[0].paragraphs[0]
    r = p.add_run(k); r.bold = True; r.font.size = Pt(10)
    meta.rows[i].cells[1].text = ''
    p = meta.rows[i].cells[1].paragraphs[0]
    r = p.add_run(v); r.font.size = Pt(10)

add_page_break(doc)

# ═════ RESUMEN EJECUTIVO ═════
add_heading(doc, 'Resumen ejecutivo', level=1)
add_para(doc,
    'Este documento consolida cinco preguntas centrales sobre la selección del motor '
    'de Reconocimiento Óptico de Caracteres (OCR) del proyecto:')
bullets = [
    '¿Cómo se construyó el Gold Standard del proyecto?',
    '¿Cuál es el propósito de ese Gold Standard?',
    '¿Qué motores OCR se evaluaron y cuáles se descartaron?',
    '¿Qué métricas cuantitativas se calcularon y qué mide cada una?',
    '¿Qué motor se eligió y con qué justificación?',
]
for b in bullets:
    doc.add_paragraph(b, style='List Number')

add_para(doc,
    'El diseño experimental se ejecutó sobre un Gold Standard de 15 documentos del corpus SECOP '
    'transcritos manualmente (36 páginas de texto humano). Se compararon dos motores OCR (EasyOCR y Tesseract) '
    'mediante cuatro métricas estándar. La decisión final se sustentó en una regla predefinida que garantiza '
    'objetividad metodológica.')

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════════════
# SECCION 1 — GOLD STANDARD
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. El Gold Standard: fundamento de la evaluación', level=1)

add_heading(doc, '1.1 ¿Qué es un Gold Standard?', level=2)
add_para(doc,
    'Para determinar si una herramienta automática funciona correctamente, '
    'se necesita un punto de referencia confiable. En Inteligencia Artificial, '
    'ese punto de referencia se denomina Gold Standard (también llamado Ground Truth '
    'o "verdad de referencia").')
add_para(doc,
    'Una analogía ayuda a entenderlo: cuando se calibra una balanza, se pesan objetos '
    'con peso conocido (pesas patrón certificadas). Esas pesas son el gold standard de '
    'la balanza. Sin ellas, no se puede responder a la pregunta "¿qué tan precisa es esta balanza?". '
    'En este proyecto, el gold standard está compuesto por 15 documentos del corpus SECOP que un '
    'anotador humano leyó y transcribió carácter por carácter. Esa transcripción humana '
    'funciona como "la verdad" contra la cual se miden los motores OCR.')

add_heading(doc, '1.2 Función del Gold Standard en las cuatro fases del proyecto', level=2)
add_para(doc,
    'El gold cumple cuatro roles durante el ciclo de vida completo del proyecto:')
add_table_from_rows(doc,
    headers=['Rol', 'Fase del proyecto', 'Uso específico'],
    rows=[
        ['Benchmark OCR', '2.1.1 — Este documento',
         'Medir la calidad de cada motor OCR contra el texto humano de referencia'],
        ['Validación de pre-anotaciones', '2.2 — Weak Supervision RUT',
         'Verificar que las reglas regex extraen las mismas entidades que el humano'],
        ['Validación del corpus completo', '2.1.4 — Cierre de gaps OCR',
         'Confirmar que las 13,254 páginas finales mantienen la calidad medida'],
        ['Evaluación de F1 del modelo NER', 'Fase 4 — Evaluación final',
         'Comparar la extracción del modelo fine-tuneado contra etiquetas humanas'],
    ],
    col_widths=[Cm(4.5), Cm(5), Cm(7)]
)
add_para(doc,
    'Sin un gold estándar, la pregunta "¿funciona este sistema?" no tiene respuesta cuantificable. '
    'Este es el principio metodológico central de CRISP-DM (Wirth & Hipp, 2000) y de toda evaluación '
    'experimental rigurosa en procesamiento de lenguaje natural (Manning, Raghavan & Schütze, 2008).')

add_heading(doc, '1.3 Diseño muestral: 15 documentos representativos', level=2)
add_para(doc,
    'La composición de los 15 documentos se definió antes de ejecutar el benchmark, '
    'bajo tres principios metodológicos:')
for b in [
    'Cobertura tipológica: representar las 4 tipologías del corpus '
    '(Cédula, RUT, Póliza, Cámara de Comercio).',
    'Estratificación por calidad visual en Cédulas: incluir tanto ejemplos nítidos como ruidosos '
    'para medir sensibilidad del OCR al ruido visual.',
    'Representatividad de escaneados: el gold contiene únicamente documentos escaneados, '
    'que son los que el OCR debe resolver (los documentos digitales van directo a PyMuPDF).',
]:
    doc.add_paragraph(b, style='List Bullet')

add_table_from_rows(doc,
    headers=['Tipología', 'Total docs', 'Criterio de selección'],
    rows=[
        ['Cédula', '6 (estratificada)', '3 alta calidad (blur_score ≥ Q3) + 3 ruidosas (blur_score ≤ Q1)'],
        ['RUT', '3', 'Escaneados (minoría dentro del RUT: 11.5% del total)'],
        ['Póliza', '3', 'Escaneadas (27% del corpus de Pólizas)'],
        ['Cámara de Comercio', '3', 'Escaneadas (9% del corpus de CC)'],
        ['TOTAL', '15 documentos (36 páginas transcritas)', ''],
    ],
    col_widths=[Cm(4), Cm(4.5), Cm(8)]
)

add_image(doc, FIGS / 'fig_gold_composicion.png', width_cm=15,
          caption='Figura 1. Composición del Gold Standard: 15 documentos distribuidos en 4 tipologías. '
                  'Las 6 Cédulas se subdividen en 3 nítidas y 3 ruidosas para capturar la variabilidad de calidad visual del corpus.')

add_heading(doc, '1.4 Metodología de construcción del Gold (5 pasos)', level=2)
add_image(doc, FIGS / 'fig_gold_proceso.png', width_cm=17,
          caption='Figura 2. Flujo metodológico de construcción del Gold Standard.')

add_para_mixed(doc, [('Paso 1. Selección reproducible. ', True, False, False),
    ('Se programó un muestreo aleatorio con semilla fija ', False, False, False),
    ('random_state=42', False, False, True),
    ('. Esto garantiza que cualquier persona que ejecute el notebook '
     'obtenga exactamente la misma selección de 15 documentos. La reproducibilidad es un '
     'requisito básico del método científico.', False, False, False)])

add_para_mixed(doc, [('Paso 2. Resolución del "mojibake" mediante índice MD5. ', True, False, False),
    ('Los nombres de archivo del SECOP presentaban caracteres alterados por conversión '
     'de codificación Windows CP-1252 → UTF-8 fallida (ejemplo: ', False, False, False),
    ('CÃÂ©dula', False, False, True),
    (' en lugar de ', False, False, False),
    ('Cédula', False, False, True),
    ('). Para evitar confusiones, se construyó un índice con el hash MD5 de cada PDF. El '
     'MD5 es una huella digital única del contenido del archivo — dos archivos con el mismo '
     'MD5 son idénticos aunque se llamen distinto. Este índice detectó además una reclasificación '
     'importante: un documento archivado como Cámara de Comercio era en realidad una Cédula.', False, False, False)])

add_para_mixed(doc, [('Paso 3. Transcripción humana literal. ', True, False, False),
    ('Para cada documento seleccionado se generó una plantilla de texto vacía. Un anotador '
     'humano (miembro del equipo) transcribió literalmente todo el texto visible, carácter por '
     'carácter, sin corregir errores del original (ortografía, acentos, abreviaciones). '
     'Se respeta el orden natural de lectura: arriba-abajo, izquierda-derecha.', False, False, False)])

add_para_mixed(doc, [('Paso 4. Validación de completitud. ', True, False, False),
    ('Un script verifica que los 15 archivos de transcripción están listos y tienen longitud '
     'razonable antes de ejecutar el benchmark. Las 15 transcripciones tienen entre 447 caracteres '
     '(Cédulas, 1 página) y 16,157 caracteres (Cámaras de Comercio, 4 páginas). Estas longitudes '
     'son coherentes con la naturaleza de cada tipo documental.', False, False, False)])

add_para_mixed(doc, [('Paso 5. Congelación inmutable. ', True, False, False),
    ('Una vez validado, el gold se marca como inmutable. Las transcripciones no se modifican, '
     'la semilla está fija, y una bandera ', False, False, False),
    ('RESAMPLE=False', False, False, True),
    (' previene re-muestreos accidentales. Esto garantiza que cualquier medición posterior contra '
     'este gold es comparable con mediciones anteriores.', False, False, False)])

add_heading(doc, '1.5 Justificación del tamaño muestral (15 documentos)', level=2)
add_para(doc,
    'La decisión de emplear una gold seed reducida (15 documentos) en lugar del gold extendido '
    'de 70 documentos originalmente planteado en el plan se sustenta en tres razones:')
for b in [
    'Costo humano proporcional al valor agregado: transcribir 70 documentos multipágina '
    'tomaría aproximadamente 50 horas de trabajo humano; 15 documentos toman ~8 horas. '
    'Para una decisión arquitectural (selección de motor OCR), 15 es suficiente.',
    'Suficiencia estadística para la decisión: 15 documentos × 2 motores = 30 puntos de datos. '
    'La varianza entre motores es lo suficientemente alta como para detectar al ganador con '
    '15 documentos (los resultados confirmaron esta suposición con diferencias de CER de hasta 0.64).',
    'Gold extendido diferido a Fase 4: el gold de 70 documentos con Cohen\'s Kappa ≥ 0.85 '
    'se construirá antes de evaluar el modelo NER, donde la precisión estadística entre '
    'anotadores sí es crítica para reportar F1 defendibles académicamente.',
]:
    doc.add_paragraph(b, style='List Bullet')

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════════════
# SECCION 2 — MOTORES OCR
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Motores OCR: selección de candidatos', level=1)

add_heading(doc, '2.1 Introducción al OCR', level=2)
add_para(doc,
    'OCR (Optical Character Recognition — Reconocimiento Óptico de Caracteres) es la tecnología '
    'que permite a una computadora interpretar texto contenido en imágenes. Cuando un documento '
    'se escanea, el sistema lo percibe como píxeles; el OCR convierte esos píxeles en caracteres '
    'procesables por software (NITs, fechas, nombres, etc.).')
add_para(doc,
    'Para este proyecto, el OCR es la única forma de recuperar información de los '
    '416 documentos escaneados del corpus SECOP (el 41% del total). Sin un motor OCR adecuado, '
    'el 93% de las Cédulas serían invisibles para el sistema NER posterior.')

add_heading(doc, '2.2 EasyOCR 1.7.2 — candidato principal', level=2)
add_para(doc,
    'EasyOCR es una librería open-source desarrollada por JaidedAI, basada en una arquitectura '
    'de dos redes neuronales profundas que operan en cascada:')
for b in [
    'CRAFT (Character Region Awareness For Text detection): detecta regiones de texto a nivel '
    'de carácter individual, no de palabra. Esto permite procesar texto en layouts atípicos, '
    'curvos o desordenados. Paper de referencia: Baek et al., CVPR 2019.',
    'CRNN (Convolutional Recurrent Neural Network): reconoce el contenido textual dentro de '
    'cada región detectada. Combina CNN (extracción de características visuales) con RNN '
    '(secuenciación de caracteres). Paper de referencia: Shi, Bai & Yao, IEEE TPAMI 2017.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para(doc, 'Características prácticas:')
for b in [
    'Soporta más de 80 idiomas, incluido español.',
    'Ejecuta en CPU (lento pero universal) o GPU (alto rendimiento).',
    'Descarga del modelo una sola vez (~60 MB para español) y operación offline posterior.',
    'Compatible con Python 3.12.',
]:
    doc.add_paragraph(b, style='List Bullet')

add_heading(doc, '2.3 Tesseract 5.5.0 — candidato comparador', level=2)
add_para(doc,
    'Tesseract OCR Engine es el motor OCR open-source más longevo en uso. Desarrollado '
    'originalmente en HP Labs (1984-1994), liberado por Google en 2005, y mantenido por la '
    'comunidad desde entonces.')
add_para(doc,
    'Desde la versión 4.0 (2018), Tesseract usa un clasificador LSTM (Long Short-Term Memory, '
    'una variante de red neuronal recurrente) para reconocer secuencias de caracteres. '
    'Paper de referencia: Smith, ICDAR 2007.')
add_para(doc, 'Características prácticas:')
for b in [
    'Requiere descarga manual del modelo de español (spa.traineddata, ~30 MB).',
    'Produce texto plano; los bounding boxes son opcionales.',
    'Muy rápido en CPU (5-10 veces más rápido que EasyOCR).',
    'Compatible con Python 3.12.',
]:
    doc.add_paragraph(b, style='List Bullet')

add_heading(doc, '2.4 Motores descartados: PaddleOCR y Donut', level=2)

add_para(doc,
    'Durante la fase de diseño del proyecto se consideraron dos motores adicionales que se '
    'descartaron antes del benchmark formal por razones técnicas bien fundamentadas.')

add_heading(doc, '2.4.1 PaddleOCR (Baidu) — descartado por incompatibilidad técnica', level=3)
add_para(doc, 'Razones del descarte:')
for b in [
    'Incompatibilidad con Python 3.12: PaddleOCR depende de PaddlePaddle, cuyo binding oficial '
    'requería (al momento del benchmark, abril 2026) versiones de Python ≤3.11. El entorno del proyecto '
    'usa Python 3.12.10, lo cual hace inviable la instalación sin un downgrade completo del entorno.',
    'Costo de migración injustificable: hacer downgrade del entorno completo para acomodar '
    'un solo motor OCR, cuando EasyOCR (con arquitectura similar basada en deep learning) '
    'es compatible, representaría trabajo innecesario.',
    'Rendimiento esperado similar a EasyOCR: según benchmarks públicos, PaddleOCR '
    'y EasyOCR tienen rendimientos comparables en español. No se esperaba una ventaja '
    'diferencial que justificara el costo de integración.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para(doc, 'Decisión: descartado sin evaluación empírica. Puede reconsiderarse si se '
         'actualiza la compatibilidad de la librería con Python 3.12.')

add_heading(doc, '2.4.2 Donut (NAVER, 2022) — descartado por arquitectura no intercambiable', level=3)
add_para(doc,
    'Donut (Document Understanding Transformer) es un modelo Visión-Lenguaje (VLM) end-to-end '
    'desarrollado por NAVER. Paper de referencia: Kim et al., ECCV 2022.')
add_para(doc,
    'Donut no es un motor OCR en sentido clásico. Recibe la imagen completa de un documento y '
    'produce directamente una estructura JSON con los campos extraídos, omitiendo la etapa de '
    'transcripción literal del texto.')
add_para(doc, 'Razones del descarte como motor OCR del proyecto:')
for b in [
    'Entrenado principalmente en inglés: Donut requeriría fine-tuning específico para documentos '
    'colombianos en español, lo cual requiere datos anotados y cómputo sustancial.',
    'Un modelo por tipo documental: Donut está diseñado para manejar un tipo de documento a la vez '
    '(un Donut entrenado para facturas, otro para recibos, etc.). Nuestro corpus tiene 4 tipologías '
    'radicalmente distintas (Cédula, RUT, Póliza, CC), lo cual requeriría entrenar 4 modelos '
    'Donut especializados.',
    'Inferencia lenta en CPU: aproximadamente 10× más lento que EasyOCR sin acceso a GPU, lo '
    'cual hace inviable el procesamiento de 1,678 páginas del corpus.',
    'Incompatibilidad con el flujo downstream: Donut produce JSON estructurado, no texto OCR. '
    'Esto rompe el pipeline posterior que espera texto plano para las Labeling Functions (§2.2), '
    'chunking (§2.3) y fine-tuning del NER (Fase 3).',
    'No es intercambiable con OCR clásico: Donut sustituiría el flujo completo OCR + LFs + NER '
    'por un modelo unificado. Es una decisión arquitectural mayor, no una elección de motor.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para(doc,
    'Decisión: Donut se descartó como motor OCR del proyecto. Se mantiene registrado como '
    'alternativa arquitectural (ALT-1) en el plan, revisitable únicamente para la tipología '
    'Cédula en Fase 4 si el F1 NER con el pipeline actual no alcanza los umbrales objetivo.')

add_heading(doc, '2.5 Tabla comparativa integral de los motores considerados', level=2)
add_para(doc,
    'La siguiente tabla consolida los 5 motores evaluados durante la fase de diseño, '
    'incluyendo los 2 que entraron al benchmark formal (EasyOCR y Tesseract), los 2 descartados '
    'antes del benchmark (PaddleOCR y Donut) y PyMuPDF (extractor complementario para PDFs digitales).')

add_table_from_rows(doc,
    headers=['Criterio', 'EasyOCR', 'Tesseract 5', 'PaddleOCR', 'Donut', 'PyMuPDF'],
    rows=[
        ['Tipo',
         'OCR deep learning', 'OCR clásico con LSTM', 'OCR deep learning',
         'VLM end-to-end', 'Extractor nativo'],
        ['Arquitectura',
         'CRAFT + CRNN', 'LSTM', 'PP-OCR',
         'Transformer Visión→JSON', 'Parser MuPDF'],
        ['Paper referencia',
         'Baek 2019 + Shi 2017', 'Smith 2007', 'Baidu 2020',
         'Kim ECCV 2022', 'Docs Artifex'],
        ['Licencia',
         'Apache 2.0', 'Apache 2.0', 'Apache 2.0',
         'MIT', 'AGPL'],
        ['Soporte español',
         'Sí (nativo)', 'Sí (spa.traineddata)', 'Sí (multi)',
         'Limitado (inglés)', 'N/A (extrae texto nativo)'],
        ['Velocidad CPU',
         'Lenta (~46 s/pág)', 'Rápida (~5 s/pág)', 'Media',
         'Muy lenta', 'Instantáneo (~0.006 s/pág)'],
        ['Entrada',
         'Imagen', 'Imagen', 'Imagen',
         'Imagen + prompt', 'PDF digital'],
        ['Salida',
         'Texto + bboxes', 'Texto plano', 'Texto + bboxes',
         'JSON directo', 'Texto plano'],
        ['Python 3.12',
         'Compatible', 'Compatible', 'NO compatible',
         'Compatible', 'Compatible'],
        ['Multi-tipología',
         'Sí (un modelo)', 'Sí', 'Sí',
         'No (un modelo/tipo)', 'Solo digitales'],
        ['Estado en el proyecto',
         'SELECCIONADO',
         'Evaluado, no principal',
         'Descartado sin evaluar',
         'Descartado sin evaluar',
         'Complementario'],
    ],
    col_widths=[Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.6), Cm(2.5), Cm(2.5)]
)

add_heading(doc, '2.6 PyMuPDF — extractor complementario (no participa del benchmark)', level=2)
add_para(doc,
    'PyMuPDF no es un motor OCR sino un extractor de texto nativo que opera sobre PDFs con '
    'caracteres digitales embebidos (no imagen). Se menciona porque es parte integral del pipeline '
    'del proyecto, aunque no participa del benchmark.')
add_para(doc,
    'En el pipeline productivo, PyMuPDF maneja los 548 documentos digitales del corpus con '
    'precisión prácticamente perfecta (CER ≈ 0), mientras que EasyOCR/Tesseract manejan los 416 '
    'escaneados. Esta bifurcación "usar el motor apropiado según el tipo de documento" es una '
    'decisión arquitectural independiente a la selección OCR.')

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════════════
# SECCION 3 — METRICAS
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Métricas de evaluación', level=1)
add_para(doc,
    'Se calcularon cuatro métricas sobre cada combinación (motor × documento). '
    'Cada métrica responde a una pregunta distinta, y el conjunto permite una evaluación '
    'multidimensional del rendimiento.')

add_heading(doc, '3.1 Character Error Rate (CER) — calidad a nivel de carácter', level=2)
add_para_mixed(doc, [('Pregunta que responde: ', True, False, False),
    ('¿qué tan cerca está el texto producido por el OCR del texto real transcrito por el humano?',
     False, True, False)])
add_para(doc, 'Fórmula:')
add_code_block(doc, 'CER = (S + D + I) / N\n\ndonde:\n  S = caracteres sustituidos (el OCR escribió un carácter incorrecto)\n  D = caracteres borrados (el OCR omitió un carácter del original)\n  I = caracteres insertados (el OCR agregó un carácter que no existía)\n  N = total de caracteres en la transcripción humana (referencia)')
add_para(doc,
    'Se calcula con la distancia de edición de Levenshtein entre el texto OCR y la transcripción '
    'humana. Es la métrica canónica de la literatura OCR desde los años 80.')
add_para_mixed(doc, [('Dirección: ', True, False, False),
    ('menor es mejor. CER = 0 significa coincidencia perfecta. CER = 0.30 significa que '
     'aproximadamente 30 de cada 100 caracteres tienen algún error.', False, False, False)])
add_para_mixed(doc, [('Normalización previa: ', True, False, False),
    ('antes de comparar se convierte todo a minúsculas y se colapsan los espacios múltiples. '
     'Esto evita que diferencias triviales de formato cuenten como errores.', False, False, False)])
add_para_mixed(doc, [('Implementación: ', True, False, False),
    ('librería jiwer — ', False, False, False),
    ('jiwer.cer(reference, hypothesis)', False, False, True)])

add_heading(doc, '3.2 Word Error Rate (WER) — calidad a nivel de palabra', level=2)
add_para_mixed(doc, [('Pregunta que responde: ', True, False, False),
    ('¿cuántas palabras enteras quedan erradas en la salida del OCR?', False, True, False)])
add_para_mixed(doc, [('Fórmula: ', True, False, False),
    ('la misma que CER pero contando palabras en lugar de caracteres. Se calcula con '
     'alineación Levenshtein sobre la secuencia de palabras tokenizadas por espacios.', False, False, False)])
add_para_mixed(doc, [('Por qué medir ambas (CER y WER): ', True, False, False),
    ('CER y WER miden dimensiones distintas. Un OCR puede tener CER bajo (pocos errores '
     'de carácter) pero WER alto si los errores se concentran en unas pocas palabras que quedan '
     'irreconocibles. La combinación permite distinguir entre "muchos errores pequeños repartidos" '
     'y "pocas palabras totalmente arruinadas".', False, False, False)])

add_heading(doc, '3.3 Entity Recall — utilidad downstream para NER', level=2)
add_para_mixed(doc, [('Pregunta que responde: ', True, False, False),
    ('¿qué tanto preserva el OCR las entidades (NITs, cédulas, fechas, montos) que el sistema '
     'necesita extraer en fases posteriores?', False, True, False)])
add_para(doc,
    'CER y WER miden la calidad del texto en general, pero lo que realmente determina el éxito '
    'de un sistema IDP es si los datos clave quedaron legibles. Un OCR con CER 0.30 puede ser '
    'perfectamente utilizable si preserva los NITs, o inútil si los NITs quedan fragmentados.')
add_para(doc, 'Fórmula:')
add_code_block(doc, 'entity_recall = |entidades_detectadas_en_OCR ∩ entidades_en_gold| / |entidades_en_gold|')
add_para(doc,
    'Las "entidades" se extraen mediante cuatro patrones regex específicos del dominio colombiano, '
    'aplicados tanto al texto OCR como a la transcripción humana de referencia:')
add_table_from_rows(doc,
    headers=['Entidad', 'Patrón regex', 'Ejemplo'],
    rows=[
        ['NIT', r'\b\d{8,10}[-\s]?\d\b', '860518862-7'],
        ['Cédula', r'\b\d{1,3}(?:[.\s]\d{3}){2,3}\b', '1.234.567 o 1 234 567'],
        ['Fecha', r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', '05/11/2025'],
        ['Monto', r'\$\s?\d{1,3}(?:[.,]\d{3})+(?:[.,]\d{1,2})?', '$1.234.567,89'],
    ],
    col_widths=[Cm(2.5), Cm(8), Cm(6)]
)
add_para_mixed(doc, [('Dirección: ', True, False, False),
    ('mayor es mejor. entity_recall = 1.00 significa que el OCR preservó todas las entidades del gold. '
     'entity_recall = 0.50 significa que la mitad de los NITs, cédulas, fechas o montos '
     'quedaron inusables para extracción posterior.', False, False, False)])
add_para_mixed(doc, [('Por qué es la métrica más importante para este proyecto: ', True, False, False),
    ('el objetivo del sistema no es "transcribir texto" sino "extraer entidades". '
     'Entity recall mide directamente esa utilidad. Esta métrica está alineada con la evaluación '
     'a nivel de entidad del shared task CoNLL-2002 (Tjong Kim Sang, 2002).', False, False, False)])

add_heading(doc, '3.4 Tiempo por página — costo operativo', level=2)
add_para_mixed(doc, [('Pregunta que responde: ', True, False, False),
    ('¿cuánto cuesta, en términos de tiempo, ejecutar este motor OCR a escala?', False, True, False)])
add_para(doc, 'Se mide con time.perf_counter() en alta resolución, contabilizando únicamente '
         'el tiempo de inferencia (no el de carga del modelo, que se amortiza).')
add_para_mixed(doc, [('Por qué es crítico: ', True, False, False),
    ('el corpus del proyecto tiene 1,678 páginas escaneadas. Un motor a 50 s/pág requiere '
     '~23 horas de cómputo (corrida overnight). Un motor a 5 s/pág requiere ~2.3 horas. La '
     'diferencia operativa es sustancial para iteración experimental.', False, False, False)])

add_heading(doc, '3.5 Resumen de las métricas calculadas', level=2)
add_table_from_rows(doc,
    headers=['Métrica', 'Qué mide', 'Unidad', 'Mejor si...', 'Pregunta que responde'],
    rows=[
        ['CER', 'Errores de carácter respecto al total', 'Proporción [0,1]', 'Menor',
         'Calidad carácter-a-carácter'],
        ['WER', 'Errores de palabra respecto al total', 'Proporción [0,1]', 'Menor',
         'Calidad palabra-a-palabra'],
        ['Entity Recall', 'Entidades preservadas en OCR', 'Proporción [0,1]', 'Mayor',
         'Utilidad para extracción NER'],
        ['Seg/página', 'Tiempo medio por página', 'Segundos', 'Menor',
         'Costo operativo a escala'],
    ],
    col_widths=[Cm(2.5), Cm(5), Cm(2.5), Cm(2), Cm(5)]
)

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════════════
# SECCION 4 — RESULTADOS COMPARATIVOS
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Comparación de métricas entre motores', level=1)

add_heading(doc, '4.1 Comparación directa de las 4 métricas (valores globales)', level=2)
add_para(doc,
    'La siguiente tabla muestra lado a lado los valores de las 4 métricas calculadas '
    'sobre los 15 documentos del Gold Standard, para cada motor. El motor ganador en cada '
    'métrica se marca con una indicación explícita.')

add_table_from_rows(doc,
    headers=['Métrica', 'EasyOCR', 'Tesseract', 'Diferencia', 'Dirección', 'Motor ganador'],
    rows=[
        ['CER (global)', '0.276', '0.446', '0.170', 'Menor es mejor', 'EasyOCR (−38%)'],
        ['WER (global)', '0.476', '0.557', '0.081', 'Menor es mejor', 'EasyOCR (−15%)'],
        ['Entity Recall', '0.551', '0.605', '0.054', 'Mayor es mejor', 'Tesseract (+10%)'],
        ['Segundos/página', '46.02', '5.06', '40.96', 'Menor es mejor', 'Tesseract (9× más rápido)'],
    ],
    col_widths=[Cm(3), Cm(2.3), Cm(2.3), Cm(2.5), Cm(3), Cm(3.9)]
)

add_image(doc, FIGS / 'fig_metricas_comparacion.png', width_cm=17,
          caption='Figura 3. Comparación visual de las 4 métricas calculadas sobre el Gold Standard. '
                  'El borde verde y la etiqueta "GANA" indican el motor vencedor en cada métrica.')

add_para(doc,
    'Lectura global:')
for b in [
    'EasyOCR gana en métricas de calidad textual (CER y WER) con márgenes significativos.',
    'Tesseract gana en entity_recall marginalmente (+10%) y en velocidad de forma decisiva (9×).',
    'Ningún motor gana en las 4 métricas simultáneamente — hay compromisos inherentes.',
]:
    doc.add_paragraph(b, style='List Bullet')

add_heading(doc, '4.2 Análisis por tipología documental (el hallazgo principal)', level=2)
add_para(doc,
    'La agregación global oculta un comportamiento crítico: cada motor domina tipologías '
    'documentales distintas. Esta tabla desglosa por tipología los resultados del CER '
    'y el entity_recall.')

add_table_from_rows(doc,
    headers=['Tipología', 'EasyOCR CER', 'Tesseract CER', 'EasyOCR ent.rec', 'Tesseract ent.rec', 'Ganador'],
    rows=[
        ['Cédula (6 docs)', '0.333', '0.782', '0.444', '0.111', 'EasyOCR (abrumador)'],
        ['RUT (3 docs)', '0.289', '0.394', '0.889', '0.889', 'Mixto'],
        ['Póliza (3 docs)', '0.329', '0.226', '0.649', '0.951', 'Tesseract'],
        ['Cámara Comercio (3 docs)', '0.096', '0.047', '0.326', '0.963', 'Tesseract (contundente)'],
    ],
    col_widths=[Cm(3.5), Cm(2.3), Cm(2.5), Cm(2.5), Cm(2.7), Cm(3.5)]
)

add_image(doc, FIGS / 'fig_comparacion_motores.png', width_cm=17,
          caption='Figura 4. Comparación por tipología documental. Izquierda: CER (menor es mejor). '
                  'Derecha: Entity Recall (mayor es mejor). Se observa claramente el régimen mixto: '
                  'EasyOCR domina Cédulas; Tesseract domina Cámara de Comercio.')

add_heading(doc, '4.3 Casos extremos y discrepancias', level=2)
add_para(doc, 'Los tres documentos con mayor diferencia de CER entre motores son todos Cédulas:')

add_table_from_rows(doc,
    headers=['Documento', 'EasyOCR CER', 'Tesseract CER', 'Diferencia'],
    rows=[
        ['23 cc.pdf', '0.261', '0.903', '0.642'],
        ['cc Julieth Payares.pdf', '0.416', '0.955', '0.540'],
        ['CC Yerlis cabarcas.pdf', '0.460', '0.998', '0.538'],
    ],
    col_widths=[Cm(6), Cm(2.5), Cm(2.5), Cm(2.5)]
)

add_para(doc,
    'Tesseract colapsa sistemáticamente en Cédulas (CER > 0.9 significa texto prácticamente '
    'irrecuperable). Este patrón confirma que la decisión de motor no puede basarse únicamente '
    'en métricas agregadas — hay que analizar el comportamiento por tipología.')

add_heading(doc, '4.4 Interpretación técnica de los resultados', level=2)

add_para_mixed(doc, [('¿Por qué Tesseract falla en Cédulas? ', True, False, False),
    ('La cédula colombiana combina condiciones adversas para un clasificador LSTM clásico: '
     'texto pequeño (6-8 pt), hologramas superpuestos, columnas apretadas y bajo contraste. '
     'Tesseract pasa la imagen completa al LSTM, que se satura. En contraste, EasyOCR detecta '
     'regiones de texto primero (mediante CRAFT) y luego reconoce cada región por separado, '
     'aislando el ruido visual. Esta estrategia de detección previa está documentada como '
     'ventajosa en Baek et al. 2019.', False, False, False)])

add_para_mixed(doc, [('¿Por qué Tesseract gana en Cámara de Comercio? ', True, False, False),
    ('Los certificados de CC son el antípoda: texto estándar de 9-10 pt, sin elementos '
     'superpuestos, columnas anchas, alto contraste. En este régimen el LSTM clásico funciona '
     'excelentemente y gana adicionalmente por velocidad (9× más rápido que EasyOCR).', False, False, False)])

add_page_break(doc)

# ═════════════════════════════════════════════════════════════════════════
# SECCION 5 — DECISION
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Decisión final y justificación', level=1)

add_image(doc, FIGS / 'fig_tradeoff.png', width_cm=14,
          caption='Figura 5. Compromiso calidad vs velocidad. El punto ideal está en la esquina '
                  'inferior-izquierda (bajo CER, baja latencia).')

add_heading(doc, '5.1 Regla de decisión preestablecida', level=2)
add_para(doc,
    'Para garantizar objetividad metodológica, la regla de decisión se definió por escrito '
    'antes de ejecutar el benchmark. Este orden previene el sesgo de adaptar la regla a los resultados.')

add_code_block(doc,
    'Regla 1: gana el motor con menor CER global, si su tiempo es menor a 2× el del motor más rápido.\n'
    'Regla 2: si hay empate en CER (diferencia < 2%), gana el motor de mayor entity_recall.\n'
    'Regla 3: si cada motor domina tipologías distintas, implementar selector híbrido.')

add_heading(doc, '5.2 Aplicación de la regla a los resultados', level=2)

add_para_mixed(doc, [('Regla 1: ', True, False, False),
    ('EasyOCR tiene menor CER global (0.276 vs 0.446), pero es 9× más lento que Tesseract. '
     '9 > 2, por tanto no cumple la restricción de tiempo. No gana por regla 1.', False, False, False)])

add_para_mixed(doc, [('Regla 2: ', True, False, False),
    ('la diferencia de CER (17 puntos absolutos) es > 2%. La regla de empate no aplica.', False, False, False)])

add_para_mixed(doc, [('Regla 3: ', True, False, False),
    ('cada motor domina tipologías distintas (EasyOCR en Cédulas; Tesseract en Cámara de Comercio y Póliza). ', False, False, False),
    ('Esta regla aplica.', True, False, False)])

add_heading(doc, '5.3 Decisión final: EasyOCR unificado', level=2)

add_para(doc,
    'Formalmente la regla 3 sugiere un selector híbrido (un motor distinto por tipología). '
    'Sin embargo, el proyecto adoptó EasyOCR unificado para todos los documentos escaneados. '
    'Las razones de esta decisión arquitectural son:')
for b in [
    'Las Cédulas son la tipología MÁS numerosa del corpus (32.9%, 334 documentos). EasyOCR es '
    'crítico para esta tipología, y su CER en las otras tipologías es aceptable (no catastrófico).',
    'Simplicidad del pipeline: un solo motor reduce puntos de falla, facilita mantenimiento '
    'y evita ramificaciones condicionales en el código productivo.',
    'Con GPU futuro: EasyOCR pasa de 46 s/pág a ~1 s/pág (40× más rápido). Esto elimina la '
    'ventaja principal de Tesseract y unifica el pipeline en un solo motor de alta calidad.',
    'Tesseract queda disponible como experimento de respaldo: si en Fase 4 el F1 del NER '
    'en Pólizas/CC queda bajo umbral, se puede reconsiderar un selector híbrido.',
]:
    doc.add_paragraph(b, style='List Bullet')

add_heading(doc, '5.4 Validación posterior en producción', level=2)
add_para(doc,
    'Para verificar que la decisión del benchmark se sostiene a escala productiva, se ejecutó '
    'EasyOCR sobre el corpus completo de 1,678 páginas escaneadas (Notebook 05, corrida overnight '
    'de 23 horas). Las métricas se midieron nuevamente contra el mismo Gold Standard.')

add_table_from_rows(doc,
    headers=['Métrica', 'Benchmark aislado (nb03)', 'Producción (nb05)', 'Variación'],
    rows=[
        ['CER global', '0.276', '0.282', '+2.2% (despreciable)'],
        ['Entity Recall', '0.551', '0.640', '+16% (mejora sustancial)'],
    ],
    col_widths=[Cm(3), Cm(4.5), Cm(4), Cm(5)]
)

add_para(doc,
    'El benchmark predijo correctamente el comportamiento productivo. La mejora de +16% en '
    'entity_recall se atribuye a la eliminación del paso binarize() del pipeline productivo — '
    'decisión técnica tomada durante la ejecución y documentada en OCR_BENCHMARK.md §2.6.0.')

add_heading(doc, '5.5 Metadatos de trazabilidad', level=2)
add_table_from_rows(doc,
    headers=['Campo', 'Valor'],
    rows=[
        ['Fecha de la decisión', '2026-04-15'],
        ['Gold seed', '15 documentos en data/gold/gold_seed_manifest.csv'],
        ['Semilla aleatoria', 'random_state=42 (reproducible)'],
        ['Cap de páginas', '4 páginas por documento'],
        ['Transcripciones congeladas', 'data/gold/transcriptions/ — inmutables desde la fecha'],
        ['EasyOCR', '1.7.2 (modelos español, modo CPU)'],
        ['Tesseract', '5.5.0 (spa.traineddata en tessdata/ local)'],
        ['Entorno', 'Python 3.12.10, Windows 11, sin GPU'],
    ],
    col_widths=[Cm(6), Cm(11)]
)

add_page_break(doc)

# ═════ FIG 11 OFICIAL ═════
add_heading(doc, '6. Visualización oficial del benchmark', level=1)
add_para(doc,
    'Figura oficial generada por el Notebook 03 al finalizar el benchmark. Izquierda: barras '
    'de CER medio por motor y tipología. Derecha: scatter con los 30 puntos de datos individuales '
    '(15 documentos × 2 motores) en el plano CER vs segundos/página.')
add_image(doc, OCR_FIG, width_cm=17,
          caption='Figura 6. Benchmark OCR oficial (fig11_ocr_benchmark.png). Este gráfico fue generado '
                  'automáticamente por el notebook de benchmark sobre los datos reales.')

add_page_break(doc)

# ═════ REFERENCIAS ═════
add_heading(doc, '7. Referencias bibliográficas', level=1)
add_para(doc,
    'Todas las referencias son verificables en línea mediante DOI, arXiv ID o URL institucional oficial.')

add_heading(doc, '7.1 Gold Standards y metodología de evaluación', level=2)
refs_a = [
    '[1] Manning, C. D., Raghavan, P., Schütze, H. (2008). Introduction to Information Retrieval. '
    'Cambridge University Press — Capítulo 8. https://nlp.stanford.edu/IR-book/',
    '[2] Wang, W. et al. (2025). A Survey on Document Intelligence Foundations and Frontiers. '
    'arXiv:2510.13366. https://arxiv.org/abs/2510.13366',
    '[3] Wirth, R. & Hipp, J. (2000). CRISP-DM: Towards a Standard Process Model for Data Mining. '
    '4th Intl. Conf. Practical Applications of Knowledge Discovery and Data Mining. '
    'https://www.cs.unibo.it/~montesi/CBD/Beatriz/10.1.1.198.5133.pdf',
    '[4] Tjong Kim Sang, E. F. (2002). Introduction to the CoNLL-2002 Shared Task: '
    'Language-Independent Named Entity Recognition. CoNLL 2002. https://aclanthology.org/W02-2024/',
]
for r in refs_a:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(0.5)

add_heading(doc, '7.2 Motores OCR evaluados', level=2)
refs_b = [
    '[5] EasyOCR — Repositorio oficial JaidedAI. https://github.com/JaidedAI/EasyOCR',
    '[6] Baek, Y., Lee, B., Han, D., Yun, S., Lee, H. (2019). Character Region Awareness for Text '
    'Detection (CRAFT). CVPR 2019. https://arxiv.org/abs/1904.01941',
    '[7] Shi, B., Bai, X., Yao, C. (2017). An End-to-End Trainable Neural Network for Image-based '
    'Sequence Recognition (CRNN). IEEE TPAMI 39(11). https://arxiv.org/abs/1507.05717',
    '[8] Tesseract OCR — Repositorio oficial. https://github.com/tesseract-ocr/tesseract',
    '[9] Smith, R. (2007). An Overview of the Tesseract OCR Engine. ICDAR 2007. '
    'https://research.google/pubs/an-overview-of-the-tesseract-ocr-engine/',
    '[10] Kim, G. et al. (2022). OCR-free Document Understanding Transformer (Donut). ECCV 2022. '
    'https://arxiv.org/abs/2111.15664',
    '[11] PyMuPDF — Documentación oficial Artifex. https://pymupdf.readthedocs.io/ — '
    'Repositorio: https://github.com/pymupdf/PyMuPDF',
]
for r in refs_b:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(0.5)

add_heading(doc, '7.3 Métricas de evaluación', level=2)
refs_c = [
    '[12] Morris, A. C., Maier, V., Green, P. (2004). From WER and RIL to MER and WIL: improved '
    'evaluation measures for connected speech recognition. Interspeech 2004. '
    'https://www.isca-speech.org/archive/interspeech_2004/morris04_interspeech.html',
    '[13] jiwer — Librería Python para cálculo de CER/WER. https://github.com/jitsi/jiwer',
    '[14] DIAN. Resolución 000110 del 11-10-2021 (estructura del RUT, fundamenta regex de NIT). '
    'https://www.dian.gov.co/normatividad/Normatividad/Resoluci%C3%B3n%20000110%20de%2011-10-2021.pdf',
]
for r in refs_c:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(0.5)

add_heading(doc, '7.4 Referencias contextuales del proyecto', level=2)
refs_d = [
    '[15] Ratner, A. et al. (2018). Snorkel: Rapid Training Data Creation with Weak Supervision. '
    'VLDB 2018. https://arxiv.org/abs/1711.10160',
    '[16] Colakoglu, G. et al. (2025). A Retrospective on Information Extraction from Documents: '
    'From Layout-aware Models to Large Language Models. arXiv:2502.18179. '
    'https://arxiv.org/abs/2502.18179',
]
for r in refs_d:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(0.5)

# ═════ APENDICE ═════
add_page_break(doc)
add_heading(doc, 'Apéndice A — Output real del notebook (tabla completa de resultados)', level=1)
add_para(doc, 'Extracto directo del output del Notebook 03 (celda 27):')
add_code_block(doc,
    '   engine  folder         n  cer_mean  wer_mean  entity_recall_mean  s_per_page_mean\n'
    '  easyocr  CAMARA DE CIO  3    0.0960    0.2229              0.3259          51.5373\n'
    '  easyocr  CEDULA         6    0.3333    0.5737              0.4444          42.0840\n'
    '  easyocr  POLIZA         3    0.3286    0.4973              0.6493          48.1290\n'
    '  easyocr  rut            3    0.2891    0.5134              0.8889          46.2573\n'
    'tesseract  CAMARA DE CIO  3    0.0469    0.0316              0.9630           5.2543\n'
    'tesseract  CEDULA         6    0.7818    0.8843              0.1111           5.4147\n'
    'tesseract  POLIZA         3    0.2256    0.3466              0.9510           4.8143\n'
    'tesseract  rut            3    0.3941    0.6402              0.8889           4.4030')

add_para_mixed(doc, [('Archivos de datos: ', True, False, False),
    ('data/processed/ocr_benchmark.csv', False, False, True),
    (' (30 filas individuales, commiteable) · ', False, False, False),
    ('data/processed/ocr_benchmark_summary.csv', False, False, True),
    (' (8 filas agregadas, commiteable) · ', False, False, False),
    ('data/processed/fig11_ocr_benchmark.png', False, False, True),
    (' (Figura 6).', False, False, False)])

add_heading(doc, 'Documentos internos del proyecto', level=2)
add_table_from_rows(doc,
    headers=['Documento', 'Ubicación'],
    rows=[
        ['Plan maestro CRISP-DM++', 'PLAN_MODELADO_CRISPDM.md'],
        ['Bitácora completa del benchmark OCR', 'OCR_BENCHMARK.md'],
        ['Propuesta de modelos (Fase 3)', 'PROPUESTA_MODELOS.md'],
        ['Reporte narrativo del benchmark', 'reports/nb03_resultados.md'],
        ['Reporte de ejecución productiva', 'reports/nb05_resultados.md'],
        ['Notebook del benchmark', 'notebooks/03_benchmark_ocr.ipynb'],
        ['Gold seed manifest', 'data/gold/gold_seed_manifest.csv'],
        ['Transcripciones humanas', 'data/gold/transcriptions/*.txt (gitignored por PII)'],
    ],
    col_widths=[Cm(7), Cm(10)]
)

# Pie
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Este documento es parte de los entregables académicos del proyecto SinergIA Lab · '
    'PUJ Especialización en IA · Ciclo 1 · 2026. Toda afirmación cuantitativa está respaldada por '
    'artefactos verificables en el repositorio del proyecto.'
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(str(OUT))
print(f'DOCX generado: {OUT}')
print(f'  Tamaño: {OUT.stat().st_size:,} bytes')
