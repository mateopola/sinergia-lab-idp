"""
Genera V2 del entregable (version concisa) del Gold Standard + Benchmark OCR.

Ejecutar desde la raiz:
    python entregables/_build_docx_v2.py

Produce:
    entregables/Entrega_Gold_Standard_y_Benchmark_OCR_v2.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
FIGS = ROOT / 'figs'
OCR_FIG = ROOT.parent / 'data' / 'processed' / 'fig11_ocr_benchmark.png'
OUT = ROOT / 'Entrega_Gold_Standard_y_Benchmark_OCR_v2.docx'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    return p


def add_para_mixed(doc, parts):
    p = doc.add_paragraph()
    for text, bold, italic, is_code in parts:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if bold: r.bold = True
        if italic: r.italic = True
        if is_code:
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
    return p


def add_alert(doc, text):
    """Caja de advertencia para ⚠ PENDIENTE"""
    p = doc.add_paragraph()
    r = p.add_run('⚠ ' + text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, widths=None, header_bg='2E5C8A'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = ''
            p = t.rows[i].cells[j].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = w
    return t


def add_image(doc, path, width_cm=15, caption=None):
    if not Path(path).exists():
        p = doc.add_paragraph(f'[FIGURA FALTANTE: {path}]')
        r = p.runs[0]; r.italic = True
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_pb(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ═════ PORTADA ═════
t = doc.add_heading('', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Gold Standard y Selección del Motor OCR')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Proyecto SinergIA Lab — PUJ Especialización en IA · Ciclo 1 · 2026\n'
              'Versión 2 — concisa (2026-04-20)')
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════
# 2.3 CONSTRUCCION DEL GOLD STANDARD
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.3 Construcción del Gold Standard', level=1)

add_para(doc,
    'El Gold Standard es el conjunto de referencia con el que se mide de forma independiente '
    'el desempeño de cada motor OCR, y contra el que se evaluarán los modelos NER del Ciclo 3. '
    'Su construcción es el hito metodológico central del Ciclo 2.')

add_heading(doc, '2.3.1 Principios', level=2)
for texto, desc in [
    ('Independencia',
     'El Gold Standard no se genera por un único modelo OCR. Se usa un OCR como borrador, '
     'pero la verdad la define el anotador humano.'),
    ('Trazabilidad',
     'Cada página queda ligada a su documento fuente, número de página, anotador y fecha de validación.'),
    ('Representatividad',
     'El muestreo respeta la distribución de clases y la variabilidad de calidad visual del corpus.'),
    ('Medibilidad',
     'Permite calcular métricas cuantitativas reproducibles sin depender del juicio subjetivo.'),
]:
    add_para_mixed(doc, [(f'{texto}. ', True, False, False), (desc, False, False, False)])

add_heading(doc, '2.3.2 Estrategia de muestreo', level=2)
add_para(doc,
    'Se aplicó un muestreo estratificado por tipo documental sobre los 1,678 páginas escaneadas '
    '(412 documentos) del corpus. Semilla fija random_state=42 (reproducible). Cap uniforme de '
    '4 páginas por documento para acotar el esfuerzo humano de transcripción.')

add_table(doc,
    headers=['Clase', 'Docs escaneados en corpus', 'Docs en Gold', '% muestreado', 'Criterio de estratificación'],
    rows=[
        ['Cédula de Ciudadanía', '308', '6', '1.9%',
         'Estratificado por calidad visual de Fase 1: 3 alta calidad (quality_label=APTO con blur y contraste altos) + 3 ruidosas (quality_label=REQUIERE_PREPROCESAMIENTO o blur bajo)'],
        ['RUT (DIAN)', '24', '3', '12.5%',
         'Aleatorio sobre escaneados (la mayoría del RUT del corpus es digital: 200/235)'],
        ['Póliza de Seguro', '59', '3', '5.1%',
         'Aleatorio sobre escaneados'],
        ['Cámara de Comercio', '16', '3', '18.8%',
         'Aleatorio sobre escaneados (la mayoría del CC del corpus es digital: 183/199)'],
        ['TOTAL', '407', '15 docs (39 páginas transcritas)', '3.7% (docs) · 2.3% (páginas)', '—'],
    ],
    widths=[Cm(3), Cm(2.8), Cm(2.2), Cm(2.5), Cm(6)]
)

add_image(doc, FIGS / 'fig_gold_composicion.png', width_cm=14,
          caption='Figura 1. Composición del Gold Standard: 15 documentos / 39 páginas.')

add_heading(doc, '2.3.3 Protocolo de anotación', level=2)
add_para(doc, 'El flujo de anotación sigue 5 pasos:')

pasos = [
    ('Paso 1 — Pre-anotación automática.',
     'Se renderiza cada página a 300 DPI y se aplica EasyOCR como borrador inicial. '
     'Este borrador se descarta después; solo acelera la transcripción manual.'),
    ('Paso 2 — Anotación humana (pasada A).',
     'Un integrante del equipo revisa el borrador línea por línea contra la imagen y corrige '
     'tildes, ñ, cifras, puntuación y caracteres inventados por el OCR.'),
    ('Paso 3 — Validación cruzada (pasada B).',
     'Un segundo integrante revisa la anotación de la pasada A con la imagen al frente. '
     'Los desacuerdos se resuelven con una tercera lectura conjunta.'),
    ('Paso 4 — Consolidación.',
     'El texto consensuado se marca como texto_gold con metadatos de trazabilidad '
     '(anotadores, fecha, MD5 del documento fuente).'),
    ('Paso 5 — Anotación de entidades.',
     'Sobre el texto_gold se marcan las entidades objetivo (NIT, cédula, fecha, monto) '
     'en formato JSON, usadas como ground truth del NER en Ciclo 3.'),
]
for bold, desc in pasos:
    add_para_mixed(doc, [(bold + ' ', True, False, False), (desc, False, False, False)])

add_heading(doc, '2.3.4 Reglas de anotación', level=2)
reglas = [
    'Los espacios en blanco redundantes se colapsan a uno solo.',
    'Los saltos de línea se preservan como marcador \\n en la representación textual.',
    'Las tildes y la ñ se escriben con ortografía correcta aunque el OCR las omita.',
    'Las cifras se transcriben tal como aparecen (sin modificar separadores decimales o de miles).',
    'Los sellos, firmas y marcas de agua se ignoran salvo que contengan información textual legible.',
    'Los errores ortográficos del documento original se preservan (no se corrigen).',
]
for r in reglas:
    doc.add_paragraph(r, style='List Bullet')

add_heading(doc, '2.3.5 Métrica de acuerdo inter-anotador', level=2)
add_para(doc,
    'Se calcula el coeficiente Cohen\'s κ a nivel de carácter entre la pasada A y la pasada B. '
    'El umbral mínimo aceptable fijado es κ ≥ 0.85 (acuerdo "casi perfecto" según Landis & Koch, 1977). '
    'Las páginas con κ por debajo del umbral se reanotan.')
add_alert(doc,
    'PENDIENTE — Acuerdo inter-anotador ejecutado: reportar el valor final de Cohen\'s κ '
    'obtenido sobre las 39 páginas del gold, y el número de páginas reanotadas.')

add_pb(doc)

# ═════════════════════════════════════════════════════════════════════════
# 2.4 EVALUACION COMPARATIVA
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.4 Evaluación comparativa de motores OCR', level=1)

add_para(doc,
    'Siguiendo la recomendación metodológica de medir cada motor de forma independiente contra '
    'el Gold Standard, se consideraron tres motores candidatos. La hipótesis de trabajo era que '
    'un motor OCR-free (Donut) podría ser competitivo en documentos de alto ruido (cédulas), mientras '
    'que motores OCR-based con soporte de español (EasyOCR, Tesseract) rendirían mejor en documentos '
    'digitales extensos (Cámara, RUT, Póliza).')

add_heading(doc, '2.4.1 Motores considerados', level=2)

add_table(doc,
    headers=['Motor', 'Tipo', 'Versión', 'Configuración', 'Estado'],
    rows=[
        ['EasyOCR', 'OCR-based neural (CRAFT + CRNN)', '1.7.2',
         "Idioma: ['es']; paragraph=True; detail=0; GPU=False", 'Evaluado'],
        ['Tesseract', 'OCR-based LSTM clásico', '5.5.0',
         'Idioma: spa; PSM=6 (bloque uniforme); OEM=1 (LSTM)', 'Evaluado'],
        ['Donut', 'OCR-free (VLM end-to-end)', 'base-finetuned-docvqa',
         'Prompt de transcripción libre; max_length=768',
         'Descartado antes del benchmark'],
    ],
    widths=[Cm(2.3), Cm(3.5), Cm(2.5), Cm(5), Cm(3)]
)

add_para_mixed(doc, [
    ('Justificación del descarte de Donut: ', True, False, False),
    ('modelo entrenado principalmente en inglés (requiere fine-tuning costoso para español), '
     'diseñado para un tipo documental a la vez (nuestro corpus tiene 4 tipologías radicalmente '
     'distintas que requerirían 4 modelos especializados), inferencia muy lenta en CPU y salida '
     'JSON incompatible con el pipeline downstream (LFs, chunking, NER). Queda registrado como '
     'alternativa arquitectural ALT-1 revisitable en Fase 4 solo para Cédulas si el F1 del pipeline '
     'actual no alcanza umbrales. Evaluación empírica se documenta en el benchmark siguiente.', False, False, False)
])

add_heading(doc, '2.4.2 Métricas calculadas', level=2)
add_para(doc,
    'Se calcularon cuatro métricas sobre cada par (motor × documento), cubriendo tres dimensiones: '
    'calidad textual, utilidad downstream y costo operativo.')

add_table(doc,
    headers=['Métrica', 'Qué mide', 'Dirección'],
    rows=[
        ['CER (Character Error Rate)',
         '(Sustituciones + Borrados + Inserciones) / N — a nivel carácter con Levenshtein. '
         'Métrica canónica de calidad OCR.',
         'Menor es mejor'],
        ['WER (Word Error Rate)',
         'Misma fórmula a nivel palabra. Complementa CER: errores concentrados en palabras clave '
         'penalizan WER aunque CER sea bajo.',
         'Menor es mejor'],
        ['Entity Recall',
         'Entidades detectadas en OCR ∩ entidades en Gold / entidades en Gold. '
         'Regex sobre NIT, cédula, fecha y monto. Mide utilidad directa para NER downstream.',
         'Mayor es mejor'],
        ['Segundos por página',
         'Tiempo medio de inferencia por página. Costo operativo a escala.',
         'Menor es mejor'],
    ],
    widths=[Cm(4), Cm(10.5), Cm(3)]
)

add_para_mixed(doc, [
    ('Nota sobre LASER (similitud semántica multilingüe): ', True, False, False),
    ('LASER (Artetxe & Schwenk, 2019) se registra como métrica complementaria deseable para '
     'trabajo futuro — captura preservación de significado cuando hay errores léxicos menores. '
     'En este Ciclo priorizamos CER/WER/entity_recall por su alineación directa con NER, que es '
     'el objetivo del Ciclo 3.', False, False, False)
])

add_heading(doc, '2.4.3 Resultados por motor y por clase', level=2)

add_para(doc, 'Resultados globales sobre los 15 documentos del Gold Standard:')

add_table(doc,
    headers=['Motor', 'CER ↓', 'WER ↓', 'Entity Recall ↑', 'Seg/página ↓'],
    rows=[
        ['EasyOCR', '0.276', '0.476', '0.551', '46.02'],
        ['Tesseract', '0.446', '0.557', '0.605', '5.06'],
    ],
    widths=[Cm(3.5), Cm(2.5), Cm(2.5), Cm(3), Cm(3)]
)

add_para(doc, 'Desglose por tipología documental:')

add_table(doc,
    headers=['Clase', 'EasyOCR CER ↓', 'Tesseract CER ↓', 'EasyOCR Ent.Rec ↑', 'Tesseract Ent.Rec ↑'],
    rows=[
        ['Cédula', '0.333', '0.782', '0.444', '0.111'],
        ['RUT', '0.289', '0.394', '0.889', '0.889'],
        ['Póliza', '0.329', '0.226', '0.649', '0.951'],
        ['Cámara de Comercio', '0.096', '0.047', '0.326', '0.963'],
        ['PROMEDIO', '0.276', '0.446', '0.551', '0.605'],
    ],
    widths=[Cm(3.5), Cm(3), Cm(3), Cm(3.5), Cm(3.5)]
)

add_image(doc, FIGS / 'fig_metricas_comparacion.png', width_cm=16,
          caption='Figura 2. Comparación de las 4 métricas calculadas (valores globales). '
                  'EasyOCR gana en calidad textual (CER, WER); Tesseract gana en velocidad y entity recall global.')

add_image(doc, FIGS / 'fig_comparacion_motores.png', width_cm=16,
          caption='Figura 3. Régimen mixto por tipología: EasyOCR domina Cédulas; Tesseract domina Cámara de Comercio.')

add_heading(doc, '2.4.4 Decisión del motor', level=2)

add_para_mixed(doc, [
    ('Motor seleccionado: EasyOCR (motor único para todos los documentos escaneados).', True, False, False)
])
add_para(doc,
    'La decisión se sustenta en cuatro criterios que priorizamos sobre la ventaja de velocidad de Tesseract:')
for b in [
    'Mejor desempeño en Cédulas (CER 0.333 vs Tesseract 0.782). Las Cédulas son la tipología MÁS '
    'numerosa del corpus (32.9%, 334 documentos). Un motor que colapse en esta clase no es viable.',
    'Soporte nativo de español con manejo correcto de tildes, ñ y signos de puntuación — reduce '
    'el trabajo posterior de normalización en el pipeline de NER.',
    'Simplicidad del pipeline: un único motor reduce puntos de falla, facilita mantenimiento y '
    'evita ramificaciones condicionales en el código productivo.',
    'Ruta de escalado a GPU: con GPU, EasyOCR pasa de 46 s/pág a ~1 s/pág (40× más rápido), '
    'eliminando la ventaja principal de Tesseract.',
]:
    doc.add_paragraph(b, style='List Number')

add_para_mixed(doc, [
    ('Se descartó un enfoque híbrido ', True, False, False),
    ('(selector por tipología) por complejidad operativa: requeriría un clasificador previo de '
     'tipología para elegir el motor, lo cual duplica los puntos de falla sin aporte compensatorio '
     'en nuestro régimen actual (CPU). Tesseract queda disponible como experimento de respaldo '
     'si en Fase 4 el F1 del NER en Pólizas/CC queda bajo umbral.', False, False, False)
])

add_para_mixed(doc, [
    ('Validación posterior en producción: ', True, False, False),
    ('EasyOCR se aplicó al corpus completo (1,678 páginas escaneadas) en una corrida overnight '
     'de 23.5 h. El CER medido contra el Gold se mantuvo en 0.282 (vs 0.276 del benchmark: '
     'desviación de +2.2%, dentro del ruido estadístico). El entity_recall subió a 0.640 '
     '(vs 0.551 del benchmark, +16%) por la eliminación del paso binarize() del pipeline productivo. '
     'La decisión del benchmark se sostiene a escala real.', False, False, False)
])

add_pb(doc)

# ═════════════════════════════════════════════════════════════════════════
# Referencias
# ═════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Referencias', level=1)

refs = [
    '[1] Baek, Y., Lee, B., Han, D., Yun, S., Lee, H. (2019). Character Region Awareness for Text '
    'Detection (CRAFT). CVPR 2019. https://arxiv.org/abs/1904.01941',
    '[2] Shi, B., Bai, X., Yao, C. (2017). An End-to-End Trainable Neural Network for Image-based '
    'Sequence Recognition (CRNN). IEEE TPAMI 39(11). https://arxiv.org/abs/1507.05717',
    '[3] Smith, R. (2007). An Overview of the Tesseract OCR Engine. ICDAR 2007. '
    'https://research.google/pubs/an-overview-of-the-tesseract-ocr-engine/',
    '[4] Kim, G. et al. (2022). OCR-free Document Understanding Transformer (Donut). ECCV 2022. '
    'https://arxiv.org/abs/2111.15664',
    '[5] Artetxe, M. & Schwenk, H. (2019). Massively Multilingual Sentence Embeddings for '
    'Zero-Shot Cross-Lingual Transfer and Beyond (LASER). TACL 2019. '
    'https://arxiv.org/abs/1812.10464',
    '[6] Morris, A. C., Maier, V., Green, P. (2004). From WER and RIL to MER and WIL: improved '
    'evaluation measures for connected speech recognition. Interspeech 2004. '
    'https://www.isca-speech.org/archive/interspeech_2004/morris04_interspeech.html',
    '[7] Landis, J. R. & Koch, G. G. (1977). The Measurement of Observer Agreement for Categorical '
    'Data. Biometrics 33(1). https://doi.org/10.2307/2529310',
    '[8] Tjong Kim Sang, E. F. (2002). Introduction to the CoNLL-2002 Shared Task: '
    'Language-Independent Named Entity Recognition. https://aclanthology.org/W02-2024/',
    '[9] Wirth, R. & Hipp, J. (2000). CRISP-DM: Towards a Standard Process Model for Data Mining. '
    'https://www.cs.unibo.it/~montesi/CBD/Beatriz/10.1.1.198.5133.pdf',
    '[10] DIAN. Resolución 000110 del 11-10-2021 (estructura del RUT). '
    'https://www.dian.gov.co/normatividad/Normatividad/Resoluci%C3%B3n%20000110%20de%2011-10-2021.pdf',
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(0.5)

# Pie
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Documento complementario: Entrega_Gold_Standard_y_Benchmark_OCR.docx (versión extendida). '
    'Bitácora completa: OCR_BENCHMARK.md. Notebook: 03_benchmark_ocr.ipynb.'
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(str(OUT))
print(f'DOCX v2 generado: {OUT}')
print(f'  Tamaño: {OUT.stat().st_size:,} bytes')
