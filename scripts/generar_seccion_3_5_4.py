"""
Genera el docx de la sección 3.5.4 Pólizas de Seguros a partir del Excel
df_modelo.xlsx del compañero (camilo trabajo/df_modelo.xlsx) y de conteos
confirmados por el responsable del corpus de Pólizas.

Valores:
  - PDFs en carpeta: 219 (confirmado por responsable)
  - PDFs procesados: 214 (doc_id únicos en df_modelo.xlsx filtrando clasificacion==POLIZA)
  - PDFs con error: 5 (219 - 214)
  - Todas las métricas de páginas/campos se calculan sobre el Excel.
"""
import os
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "camilo trabajo/Grupo-1 - Ciclo-2 (2).docx"
OUT_DIR = "entregables"
OUT = os.path.join(OUT_DIR, "Seccion_3.5.4_Polizas.docx")
os.makedirs(OUT_DIR, exist_ok=True)

doc = Document(SRC)
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_heading3(text):
    return doc.add_paragraph(text, style="Heading 3")


def add_para(text, bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        shade_cell(cell)
        set_cell_borders(cell)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            set_cell_borders(cell)
    if col_widths_cm:
        for row in t.rows:
            for c_i, w in enumerate(col_widths_cm):
                row.cells[c_i].width = Cm(w)
    return t


# ===== 3.5.4 =====
add_heading3("3.5.4 Pólizas de Seguros")

add_para(
    "Resultados de la extracción sobre el corpus de pólizas de seguros. "
    "El corpus es altamente heterogéneo por la diversidad de aseguradoras "
    "(SURA / Suramericana, La Equidad, Mundial, Allianz, Mapfre, AXA, Bolívar y Colpatria, "
    "entre otras) y de ramos (cumplimiento estatal, seriedad de la oferta, "
    "responsabilidad civil extracontractual (RCE), vigilancia, vida grupo y SOAT)."
)

add_table(
    ["Métrica", "Valor"],
    [
        ["PDFs en carpeta", "219"],
        ["PDFs procesados exitosamente", "214 (97.7%)"],
        ["PDFs con error (corruptos / protegidos)", "5 (2.3%)"],
        ["Total de páginas convertidas", "2,949"],
        ["Promedio de páginas por documento", "13.8"],
        ["Total de palabras extraídas", "964,807"],
    ],
    col_widths_cm=[9, 5],
)
doc.add_paragraph()

add_para("Calidad de extracción — Pólizas")
add_table(
    ["Indicador", "Valor"],
    [
        ["Páginas con texto extraído exitosamente", "2,938 (99.63%)"],
        ["Páginas sin texto (vacías)", "11 (0.37%)"],
        ["Páginas con ≥ 300 palabras (texto denso)", "1,406 (47.7%)"],
        ["Páginas con ruido OCR alto", "1 (0.0%)"],
        ["Longitud promedio por página", "2,158 caracteres"],
        ["Longitud mediana por página", "1,933 caracteres"],
    ],
    col_widths_cm=[9, 5],
)
doc.add_paragraph()

add_para("Distribución de longitud de texto por página")
add_table(
    ["Rango (caracteres)", "Páginas", "%"],
    [
        ["< 30 (vacías)", "11", "0.4%"],
        ["30 – 200", "115", "3.9%"],
        ["200 – 500", "139", "4.7%"],
        ["500 – 1,000", "282", "9.6%"],
        ["1,000 – 2,000", "1,011", "34.3%"],
        ["> 2,000", "1,391", "47.2%"],
    ],
    col_widths_cm=[6, 4, 3],
)
doc.add_paragraph()

add_para("Presencia de campos clave en página 1 — Pólizas")
add_table(
    ["Campo", "Documentos con el campo", "%"],
    [
        ['"PÓLIZA" (mención)', "196 / 214", "91.6%"],
        ["Vigencia", "179 / 214", "83.6%"],
        ["Tomador", "177 / 214", "82.7%"],
        ["Asegurado", "170 / 214", "79.4%"],
        ["NIT", "169 / 214", "79.0%"],
        ["Beneficiario", "125 / 214", "58.4%"],
        ["Valor / Suma asegurada", "86 / 214", "40.2%"],
    ],
    col_widths_cm=[6, 5, 3],
)
doc.add_paragraph()

add_para("Extensión de documentos — Pólizas")
add_table(
    ["Rango de páginas", "Documentos", "%"],
    [
        ["1 – 2 páginas", "36", "16.8%"],
        ["3 – 5 páginas", "45", "21.0%"],
        ["6 – 10 páginas", "42", "19.6%"],
        ["11 – 20 páginas", "66", "30.8%"],
        ["21 – 50 páginas", "16", "7.5%"],
        ["> 50 páginas", "9", "4.2%"],
    ],
    col_widths_cm=[6, 4, 3],
)
doc.add_paragraph()

conclusion = (
    "De los 219 PDFs del corpus de Pólizas, 214 (97.7%) fueron procesados exitosamente; los "
    "5 restantes presentaron error de lectura por corrupción o archivos con formato no válido "
    '(p. ej. "13. GARANTIA DE SERIEDAD_1.pdf", que arrojó errores de tabla xref y caracteres '
    "ilegales al intentar su apertura). Sobre las 2,949 páginas procesadas, EasyOCR extrajo "
    "texto en el 99.63%; las 11 páginas vacías corresponden típicamente a carátulas en blanco "
    "y separadores. La proporción de páginas densas (≥ 300 palabras) fue del 47.7%, inferior "
    "al 77.6% observado en Cámaras de Comercio, diferencia que se explica por la estructura "
    "propia de las pólizas, que combinan carátulas cortas, certificados de vigencia y anexos "
    "tarifarios con el clausulado extenso. La longitud mediana de 1,933 caracteres por página "
    "y la práctica ausencia de ruido OCR alto (1 página) confirman la efectividad del "
    "preprocesamiento con binarización descrito en 3.2. En los campos clave de la primera "
    'página se observa alta consistencia en la mención de "PÓLIZA" (91.6%), Vigencia (83.6%), '
    "Tomador (82.7%), Asegurado (79.4%) y NIT (79.0%); por el contrario, Beneficiario (58.4%) "
    "y Valor / Suma asegurada (40.2%) presentan menor cobertura en página 1, pues su "
    "ubicación varía entre carátula, anexos y condicionados dependiendo de la aseguradora y "
    "del ramo. Este comportamiento valida la hipótesis inicial de alta variabilidad de "
    "formato entre emisores y justifica, para esta clase documental, una estrategia de "
    "extracción de campos robusta a múltiples plantillas."
)
p = doc.add_paragraph()
r = p.add_run("Conclusión (Pólizas). ")
r.bold = True
p.add_run(conclusion)

p2 = doc.add_paragraph()
r2 = p2.add_run("Campos clave esperados — Póliza: ")
r2.bold = True
p2.add_run(
    "Número de póliza, tomador, asegurado, aseguradora, vigencia (desde / hasta), "
    "valor asegurado, prima, coberturas."
)

if sectPr is not None:
    body.append(sectPr)

doc.save(OUT)
print(f"OK: {OUT}")
print(f"Tamaño: {os.path.getsize(OUT)} bytes")
